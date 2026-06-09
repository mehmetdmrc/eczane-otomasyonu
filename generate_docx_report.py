import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    tcPr.append(parse_xml(shd_xml))

def create_report():
    doc = docx.Document()
    
    # Sayfa kenar boşlukları
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Standart Siyah/Beyaz Renk Paleti
    BLACK_COLOR = RGBColor(0, 0, 0)
    WHITE_COLOR = RGBColor(255, 255, 255)

    # Başlık Stili
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("ECZANE BİLGİ YÖNETİM SİSTEMİ (EBYS)\nPROJE RAPORU")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = BLACK_COLOR
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Teknik, Mimari ve İşlevsel İnceleme Raporu")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = BLACK_COLOR
    
    doc.add_paragraph() # Boşluk
    
    # İÇİNDEKİLER BÖLÜMÜ
    h_toc = doc.add_heading("İÇİNDEKİLER", level=1)
    h_toc.runs[0].font.color.rgb = BLACK_COLOR
    
    toc_items = [
        "1. Projenin Amacı",
        "2. Kullanılan Teknolojiler",
        "    • İstemci (Kullanıcı Arayüzü)",
        "    • Sunucu (Backend & REST API)",
        "    • Veritabanı (Database)",
        "3. Sistem Modülleri",
        "    • Yetkilendirme ve Ana Sayfa Modülü",
        "    • İlaç ve Stok Yönetimi Modülü",
        "    • Hasta Yönetimi Modülü",
        "    • Satış Yönetimi Modülü",
        "    • Raporlama Modülü",
        "    • Ayarlar Modülü",
        "4. Veritabanı Yapısı",
        "5. Öne Çıkan Özellikler ve Kurulum Detayları"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK_COLOR
        if not item.startswith("    "):
            run.font.bold = True
            p.paragraph_format.space_before = Pt(6)
        else:
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # 1. PROJENİN AMACI
    h1 = doc.add_heading("1. Projenin Amacı", level=1)
    h1.runs[0].font.color.rgb = BLACK_COLOR
    
    p = doc.add_paragraph()
    run = p.add_run(
        "Eczane Bilgi Yönetim Sistemi (EBYS), modern eczanelerin günlük iş süreçlerini (ilaç satışı, stok takibi, "
        "hasta kayıtları, finansal raporlama ve kullanıcı yetkilendirme) dijitalleştirerek hata payını en aza indiren, "
        "hızlı, güvenli ve verimli bir masaüstü otomasyon yazılımıdır."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK_COLOR
    
    p = doc.add_paragraph()
    run = p.add_run("Projenin temel amacı şunlardır:")
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK_COLOR
    
    objectives = [
        "Eczane çalışanlarının (müdür ve personeller) saniyeler içinde ilaç sorgulaması ve reçeteli/reçetesiz satış yapabilmesini sağlamak.",
        "Kritik ilaç stoklarının anlık olarak izlenmesini mümkün kılarak olası stok tükenmelerinin önüne geçmek.",
        "Hasta bilgilerinin ve geçmiş alışverişlerinin düzenli bir şekilde saklanmasıyla müşteri memnuniyetini ve takip kolaylığını artırmak.",
        "Elde edilen veriler üzerinden günlük, haftalık ve aylık analizler sunarak işletme kararlarının veri odaklı alınmasını kolaylaştırmak."
    ]
    for obj in objectives:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(obj)
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK_COLOR
    
    doc.add_paragraph()

    # 2. KULLANILAN TEKNOLOJİLER
    h2 = doc.add_heading("2. Kullanılan Teknolojiler", level=1)
    h2.runs[0].font.color.rgb = BLACK_COLOR
    
    p = doc.add_paragraph()
    run = p.add_run("Proje, güncel yazılım mühendisliği prensiplerine uygun olarak 3 katmanlı (İstemci - Sunucu API - Veritabanı) modern bir istemci-sunucu mimarisi üzerine inşa edilmiştir.")
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK_COLOR
    
    # İstemci
    h_sub = doc.add_heading("• İstemci (Kullanıcı Arayüzü)", level=2)
    h_sub.runs[0].font.color.rgb = BLACK_COLOR
    client_tech = [
        ("Dil & Platform: ", "C# / .NET Framework (Windows Forms)"),
        ("Mimari: ", "Modüler ve form tabanlı arayüz yapısı."),
        ("İletişim Katmanı: ", "ApiService.cs sınıfı aracılığıyla arka plan servislerine asenkron/senkron HTTP istekleri gönderilir. Veri alışverişi standart JSON formatında gerçekleştirilir."),
        ("Esneklik: ", "api_config.txt yapılandırma dosyası sayesinde API sunucu adresi dinamik olarak yönetilir; böylece yerel veya uzak sunucu geçişleri kaynak kodu değiştirmeden anında yapılabilir.")
    ]
    for b_text, n_text in client_tech:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(b_text)
        r1.bold = True
        r1.font.color.rgb = BLACK_COLOR
        r2 = p.add_run(n_text)
        r2.font.color.rgb = BLACK_COLOR
        
    # Sunucu
    h_sub = doc.add_heading("• Sunucu (Backend & REST API)", level=2)
    h_sub.runs[0].font.color.rgb = BLACK_COLOR
    server_tech = [
        ("Dil & Ortam: ", "PHP 8.x (XAMPP / Apache web sunucusu üzerinde çalışır)."),
        ("API Mimarisi: ", "Tek noktadan yönetim sağlayan RESTful API (API.php)."),
        ("Veri İletişim Protokolü: ", "HTTP metodlarına (GET, POST, PUT, DELETE, OPTIONS) tam uyumlu yönlendirme (routing) sistemi."),
        ("Güvenlik & Standartlar: ", "Access-Control-Allow-Origin (CORS) yapılandırmaları ile dış erişim güvenliği, tam UTF-8 karakter desteği ve SQL Injection saldırılarına karşı mysqli_real_escape_string ile güvenli parametre işleme.")
    ]
    for b_text, n_text in server_tech:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(b_text)
        r1.bold = True
        r1.font.color.rgb = BLACK_COLOR
        r2 = p.add_run(n_text)
        r2.font.color.rgb = BLACK_COLOR

    # Veritabanı
    h_sub = doc.add_heading("• Veritabanı (Database)", level=2)
    h_sub.runs[0].font.color.rgb = BLACK_COLOR
    db_tech = [
        ("Veritabanı Yönetim Sistemi: ", "MySQL (eczane_db)."),
        ("Bağlantı Sürücüsü: ", "PHP mysqli uzantısı."),
        ("Veri Tutarlılığı (ACID): ", "Satış ve stok güncellemelerinin aynı anda eksiksiz gerçekleşmesi için gelişmiş veritabanı işlemleri (mysqli_begin_transaction, commit ve hata anında rollback) kullanılmaktadır.")
    ]
    for b_text, n_text in db_tech:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(b_text)
        r1.bold = True
        r1.font.color.rgb = BLACK_COLOR
        r2 = p.add_run(n_text)
        r2.font.color.rgb = BLACK_COLOR

    doc.add_paragraph()

    # 3. SİSTEM MODÜLLERİ
    h3 = doc.add_heading("3. Sistem Modülleri", level=1)
    h3.runs[0].font.color.rgb = BLACK_COLOR
    
    p = doc.add_paragraph()
    p.add_run("Sistem, eczanenin tüm operasyonel ihtiyaçlarını karşılamak üzere 6 temel modüle ayrılmıştır:").font.color.rgb = BLACK_COLOR
    
    modules = [
        ("Yetkilendirme ve Ana Sayfa Modülü", [
            ("Giriş Ekranı (LoginForm.cs): ", "Kullanıcı adı ve şifre doğrulaması yaparak sisteme güvenli giriş sağlar. Arka planda user/login API uç noktasını çağırır."),
            ("Rol Bazlı Yetkilendirme: ", "Sistemde Müdür (Admin) ve Personel olmak üzere farklı yetki seviyeleri mevcuttur. Raporlama ve ayarlar gibi hassas modüllere erişim rol bazlı olarak kısıtlanır."),
            ("Ana Sayfa (formanasayfa.cs): ", "Giriş yapıldıktan sonra tüm sistem modüllerine tek tıkla ulaşımı sağlayan, özet istatistiklerin yer aldığı merkezi yönetim panelidir.")
        ]),
        ("İlaç ve Stok Yönetimi Modülü", [
            ("İlaç Ekleme ve Düzenleme (IlacEkleForm.cs): ", "İlaç adı, kategori, fiyat, stok miktarı ve kullanım açıklaması bilgileriyle yeni ilaç kaydı oluşturur veya mevcut ilaçları günceller (ilac REST uç noktası)."),
            ("Stok İşlemleri ve İzleme (stokıslemlerı.cs, IlacRaporForm.cs): ", "İlaçların stok seviyelerini listeler. Kritik stok seviyesine yaklaşan ilaçlar için kullanıcıyı bilgilendirir, depo sipariş süreçlerine altyapı oluşturur.")
        ]),
        ("Hasta Yönetimi Modülü", [
            ("Kayıt ve Profil (HastaEkleForm.cs): ", "Hastaların TC Kimlik Numarası, ad, soyad, telefon ve açık adres bilgilerini sisteme kaydeder (hasta API rotası)."),
            ("Hasta Listeleme ve Arama (HastaRaporForm.cs): ", "Hastaları listeler, TC veya isim bazlı arama yapılmasını sağlayarak satış anında hasta seçimini hızlandırır.")
        ]),
        ("Satış Yönetimi Modülü", [
            ("Satış Ekranı (satisyonetimi.cs): ", "Reçeteli veya reçetesiz işlemler için sepet mantığıyla çalışan dinamik satış modülüdür."),
            ("İşlem Akışı: ", "İlaçlar seçilip sepete eklenir. Hasta seçimi yapılır ve işlem onaylandığında satis rotasına POST isteği atılır. Veritabanında işlem eş zamanlı olarak yazılır ve stoktan anında düşülür.")
        ]),
        ("Raporlama Modülü", [
            ("Haftalık ve Aylık Satış Raporları (HaftalikRaporForm.cs, AylikSatisRaporForm.cs): ", "Belirli tarih aralıklarındaki toplam ciro ve satış adetlerini analiz eder."),
            ("En Çok ve En Az Satan İlaçlar (EnCokSatanlarForm.cs, EnAzSatanlarForm.cs): ", "Stok sirkülasyonunu izleyerek çok talep gören ürünlerin tedarik edilmesini, satılmayan ürünlerin ise tespitini sağlar."),
            ("Müşteri Analizi (EnCokAlisverisYapanlarForm.cs): ", "Eczaneden en sık alışveriş yapan ve ciro sağlayan hastaları sıralar.")
        ]),
        ("Ayarlar Modülü", [
            ("Sistem Yapılandırması (AyarlarForm.cs): ", "İstemcinin bağlandığı API sunucu adresini (api_config.txt) uygulama içerisinden görsel bir arayüzle değiştirme imkanı sunar. Böylece ağ değişikliklerinde teknik desteğe ihtiyaç duyulmaz.")
        ])
    ]
    
    for mod_title, mod_items in modules:
        mh = doc.add_heading(f"• {mod_title}", level=2)
        mh.runs[0].font.color.rgb = BLACK_COLOR
        for b_t, n_t in mod_items:
            p = doc.add_paragraph(style='List Bullet')
            r1 = p.add_run(b_t)
            r1.bold = True
            r1.font.color.rgb = BLACK_COLOR
            r2 = p.add_run(n_t)
            r2.font.color.rgb = BLACK_COLOR

    doc.add_page_break()

    # 4. VERİTABANI YAPISI
    h4 = doc.add_heading("4. Veritabanı Yapısı", level=1)
    h4.runs[0].font.color.rgb = BLACK_COLOR
    
    p = doc.add_paragraph()
    p.add_run("Sistem, ilişkisel veritabanı (RDBMS) standartlarına uygun olarak tasarlanmış 6 ana tablodan oluşmaktadır. Aşağıda veritabanı tablolarının genel yapısı özetlenmiştir:").font.color.rgb = BLACK_COLOR
    
    # Tablo Oluşturma
    table_data = [
        ["Tablo Adı", "Önemli Alanlar", "Açıklama / İlişki"],
        ["kullanicilar", "id (PK), kullanici_adi, sifre, rol", "Sistem kullanıcılarını ve rollerini (Müdür/Personel) tutar."],
        ["kategoriler", "id (PK), ad", "İlaç kategorilerini listeler."],
        ["ilaclar", "id (PK), ad, kategori_id (FK), stok, fiyat", "İlaç ve depo stok bilgilerini barındırır. Kategoriler tablosuna bağlıdır."],
        ["hastalar", "id (PK), tc (Unique), ad, soyad, telefon", "Hasta ve müşteri iletişim/kimlik bilgilerini saklar."],
        ["satislar", "id (PK), tarih, kullanici_id, hasta_id, tutar", "Satış fişi başlık verileridir. Kullanıcı ve Hasta tablolarıyla ilişkilidir."],
        ["satis_detay", "id (PK), satis_id, ilac_id, adet, birim_fiyat", "Fiş içerisindeki satır kalemlerini tutar. İlaçlar ve Satışlar ile ilişkilidir."]
    ]
    
    table = doc.add_table(rows=len(table_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Sütun genişlikleri
    col_widths = [Inches(1.5), Inches(2.5), Inches(2.5)]
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    # Başlık Hücre Biçimlendirme (Sade Gri / Siyah Konsept)
    for col_idx, text in enumerate(table_data[0]):
        cell = table.rows[0].cells[col_idx]
        set_cell_background(cell, "333333")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = WHITE_COLOR
        
    # İçerik Hücre Biçimlendirme
    for row_idx, row_content in enumerate(table_data[1:], start=1):
        bg_color = "F8F9FA" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_content):
            cell = table.rows[row_idx].cells[col_idx]
            set_cell_background(cell, bg_color)
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.05)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK_COLOR

    doc.add_paragraph() # Boşluk

    # 5. ÖNE ÇIKAN ÖZELLİKLER VE KURULUM DETAYLARI
    h5 = doc.add_heading("5. Öne Çıkan Özellikler ve Kurulum Detayları", level=1)
    h5.runs[0].font.color.rgb = BLACK_COLOR
    
    sh1 = doc.add_heading("• Öne Çıkan Sistem Özellikleri", level=2)
    sh1.runs[0].font.color.rgb = BLACK_COLOR
    
    features = [
        ("İzole ve Modüler Servis Mimarisi: ", "İstemci ile veritabanı arasında doğrudan SQL bağlantısı yerine PHP REST API (API.php) kullanılması sistemin güvenliğini üst düzeye taşır ve gelecekte Mobil veya Web arayüzü entegrasyonunu kolaylaştırır."),
        ("Kusursuz Veri Bütünlüğü (ACID Transactions): ", "Çoklu ilaç içeren satış işlemlerinde, veritabanı tarafında START TRANSACTION ile başlatılan süreç, ancak tüm satırlar başarıyla eklendiğinde ve stoklar güncellendiğinde onaylanır (COMMIT). Hata anında tüm işlemler geri alınır (ROLLBACK)."),
        ("Dinamik API Yönlendirme: ", "api_config.txt tabanlı altyapı sayesinde, C# uygulamasının yeniden derlenmesine gerek kalmaksızın sunucu adresi ve port güncellemeleri anında yapılabilir."),
        ("Zengin İstatistik ve Raporlama: ", "Eczane yöneticisine ürün sirkülasyonu, en aktif hastalar ve personel performansı hakkında çok yönlü karar destek analizleri sunar.")
    ]
    for b_t, n_t in features:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(b_t)
        r1.bold = True
        r1.font.color.rgb = BLACK_COLOR
        r2 = p.add_run(n_t)
        r2.font.color.rgb = BLACK_COLOR
        
    sh2 = doc.add_heading("• Kurulum Detayları ve Çalıştırma Adımları", level=2)
    sh2.runs[0].font.color.rgb = BLACK_COLOR
    
    p = doc.add_paragraph()
    r = p.add_run("1. Sunucu ve Veritabanı Kurulumu")
    r.bold = True
    r.font.color.rgb = BLACK_COLOR
    
    server_install = [
        "Sisteminize XAMPP (veya benzeri bir Apache/MySQL dağıtımı) kurun ve Apache/MySQL servislerini başlatın.",
        "Proje dizininde yer alan db_setup.sql dosyasını phpMyAdmin üzerinden veya MySQL komut satırından içe aktararak eczane_db veritabanını ve örnek verileri oluşturun.",
        "API.php ve proje dosyalarının C:\\xampp\\htdocs\\eczanesyp klasöründe yer aldığından emin olun.",
        "Tarayıcınızda http://localhost/eczanesyp/API.php/ilac rotasını test ederek verilerin ulaştığını doğrulayın."
    ]
    for st in server_install:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.4)
        r = p.add_run(st)
        r.font.color.rgb = BLACK_COLOR
        r.font.size = Pt(10.5)
        
    p = doc.add_paragraph()
    r = p.add_run("2. İstemci Uygulaması Kurulumu")
    r.bold = True
    r.font.color.rgb = BLACK_COLOR
    
    client_install = [
        "otomosyan projesi.sln çözüm dosyasını Visual Studio ile açın ve NuGet paketlerini geri yükleyin.",
        "bin\\Debug veya ana dizindeki api_config.txt dosyasının içeriğinde doğru API adresinin yazılı olduğunu kontrol edin.",
        "Projeyi derleyip çalıştırın (F5).",
        "Açılan giriş ekranında varsayılan yönetici hesabı olan Kullanıcı Adı: admin, Şifre: admin123 bilgileriyle sisteme giriş yapabilirsiniz."
    ]
    for ct in client_install:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.4)
        r = p.add_run(ct)
        r.font.color.rgb = BLACK_COLOR
        r.font.size = Pt(10.5)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run("Bu rapor, Eczane Bilgi Yönetim Sistemi (EBYS) projesinin teknik, mimari ve işlevsel detaylarını belgelemek amacıyla hazırlanmıştır.")
    f_run.font.size = Pt(9.5)
    f_run.font.italic = True
    f_run.font.color.rgb = BLACK_COLOR

    output_path = os.path.join(os.getcwd(), "syjavarapor.docx")
    doc.save(output_path)
    print(f"Rapor basariyla {output_path} konumuna kaydedildi.")

if __name__ == "__main__":
    create_report()
